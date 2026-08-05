"""Import the 38 biology knowledge-list DOCX files into VitePress pages.

Usage:
    python tools/import_biology_docx.py <source-directory>

The importer intentionally stops before the exercise section. It keeps the
source documents read-only and writes only the biology Markdown/assets tree.
"""

from __future__ import annotations

import io
import re
import shutil
import sys
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
BIOLOGY_DIR = ROOT / "生物"
ASSET_DIR = ROOT / "public" / "biology"


@dataclass(frozen=True)
class Module:
    name: str
    emoji: str
    start: int
    end: int
    description: str


MODULES = (
    Module("分子与细胞", "🔬", 1, 9, "从生命系统的基本单位出发，掌握细胞的物质基础、结构、代谢与生命历程。"),
    Module("遗传与进化", "🧬", 10, 18, "串联减数分裂、遗传规律、分子遗传学、生物变异与进化。"),
    Module("稳态与调节", "🫀", 19, 26, "理解动物和植物生命活动的调节机制，以及内环境稳态的维持。"),
    Module("生物与环境", "🌿", 27, 31, "从种群、群落到生态系统，建立生态学的层级分析框架。"),
    Module("生物技术", "🧫", 32, 38, "掌握发酵、细胞工程、胚胎工程和基因工程的原理、流程及安全伦理。"),
)


def module_for(number: int) -> Module:
    return next(module for module in MODULES if module.start <= number <= module.end)


def clean(text: str) -> str:
    text = text.replace("\u00a0", " ").replace("\u3000", " ").replace("_", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\s*([，。；：！？])\s*", r"\1", text)
    text = text.strip()
    if len(text) % 2 == 0 and text[: len(text) // 2] == text[len(text) // 2 :]:
        text = text[: len(text) // 2]
    return text


def iter_blocks(document: DocxDocument):
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def paragraph_lines(paragraph: Paragraph) -> list[str]:
    # paragraph.text omits text stored in Word drawing/text-box nodes. Reading
    # all w:t nodes keeps the visible headings used by these source documents.
    value = clean("".join(node.text or "" for node in paragraph._p.xpath(".//w:t")))
    return [value] if value else []


def table_lines(table: Table) -> list[str]:
    lines: list[str] = []
    seen: set[str] = set()
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for value in paragraph_lines(paragraph):
                    if value and value not in seen:
                        seen.add(value)
                        lines.append(value)
    return lines


def unique_lines(lines: list[str]) -> list[str]:
    result: list[str] = []
    for line in lines:
        if not line or line in result:
            continue
        result.append(line)
    return result


def normalize_list(lines: list[str]) -> list[str]:
    result: list[str] = []
    for line in unique_lines(lines):
        line = re.sub(r"^知识清单\d+\s*", "", line)
        if not line:
            continue
        result.append(line)
    return result


def parse_topic(path: Path):
    document = Document(path)
    match = re.search(r"知识清单(\d{2})\s+(.+?)（", path.stem)
    if not match:
        raise ValueError(f"Cannot parse topic filename: {path.name}")
    number = int(match.group(1))
    title = match.group(2).strip()

    points: list[tuple[str, list[str]]] = []
    mistakes: list[tuple[str, list[str]]] = []
    current_kind: str | None = None
    current_title = ""
    current_lines: list[str] = []

    def flush() -> None:
        nonlocal current_lines
        if current_kind == "point" and current_title:
            points.append((current_title, normalize_list(current_lines)))
        elif current_kind == "mistake" and current_title:
            content = normalize_list(current_lines)
            number_match = re.match(r"易错点(\d+)", current_title)
            previous_match = re.match(r"易错点(\d+)", mistakes[-1][0]) if mistakes else None
            if number_match and previous_match and number_match.group(1) == previous_match.group(1):
                old_title, old_content = mistakes[-1]
                extra_title = re.sub(r"^易错点\d+\s*", "", current_title)
                mistakes[-1] = (f"{old_title}；{extra_title}", unique_lines(old_content + content))
            else:
                mistakes.append((current_title, content))
        current_lines = []

    for block in iter_blocks(document):
        lines = paragraph_lines(block) if isinstance(block, Paragraph) else table_lines(block)
        for line in lines:
            if re.match(r"^1[．.]", line) and "高考真题" in line:
                flush()
                return number, title, points, mistakes
            heading = re.match(r"^(考点|易错点)\s*(\d+)\s*(.*)$", line)
            if heading:
                flush()
                current_kind = "point" if heading.group(1) == "考点" else "mistake"
                current_title = clean(f"{heading.group(1)}{heading.group(2)} {heading.group(3)}")
                continue
            if current_kind:
                current_lines.append(line)
    flush()
    return number, title, points, mistakes


def image_extension(blob: bytes, fallback: str = ".png") -> str:
    try:
        fmt = Image.open(io.BytesIO(blob)).format
        return {"JPEG": ".jpg", "PNG": ".png", "GIF": ".gif", "WEBP": ".webp"}.get(fmt, fallback)
    except Exception:
        return fallback


def extract_maps(path: Path) -> dict[int, str]:
    document = Document(path)
    maps: dict[int, tuple[int, bytes, str]] = {}
    current: int | None = None
    for paragraph in document.paragraphs:
        text = clean(paragraph.text)
        match = re.match(r"^(\d{2})\s*", text)
        if match and 1 <= int(match.group(1)) <= 38:
            current = int(match.group(1))
        if current is None:
            continue
        for blip in paragraph._p.xpath(".//a:blip"):
            rel_id = blip.get(qn("r:embed"))
            if not rel_id:
                continue
            part = document.part.related_parts[rel_id]
            blob = part.blob
            try:
                with Image.open(io.BytesIO(blob)) as image:
                    area = image.width * image.height
            except Exception:
                area = len(blob)
            # Each numbered heading is immediately followed by its map. Later
            # images can include promotional end matter, so retain the first.
            if current not in maps:
                maps[current] = (area, blob, image_extension(blob, Path(part.partname).suffix))
    if len(maps) != 38:
        missing = sorted(set(range(1, 39)) - set(maps))
        raise RuntimeError(f"Expected 38 topic maps; missing: {missing}")

    outputs: dict[int, str] = {}
    for number, (_, blob, suffix) in maps.items():
        topic_dir = ASSET_DIR / f"topic-{number:02d}"
        topic_dir.mkdir(parents=True, exist_ok=True)
        output = topic_dir / f"mind-map{suffix}"
        output.write_bytes(blob)
        outputs[number] = output.name
    return outputs


def markdown_bullets(lines: list[str]) -> list[str]:
    output: list[str] = []
    for line in lines:
        marker = re.match(r"^[（(](\d+)[）)]\s*(.*)$", line)
        if marker:
            output.append(f"- **{marker.group(1)}.** {marker.group(2)}")
        else:
            output.append(f"- {line}")
    return output or ["- 结合教材图示，理解本考点的概念、条件和结论。"]


def write_topic(topic, map_filename: str, previous, following) -> None:
    number, title, points, mistakes = topic
    module = module_for(number)
    lines = [
        f"# {number:02d} {title}",
        "",
        f"> 所属模块：[{module.name}](./) · 本页整理 {len(points)} 个核心考点和 {len(mistakes)} 个易错点。",
        "",
        "## 考点概览",
        "",
    ]
    lines.extend(f"- {heading}" for heading, _ in points)
    lines.extend(["", "## 核心知识", ""])
    for heading, content in points:
        lines.extend([f"### {heading}", ""])
        lines.extend(markdown_bullets(content))
        lines.append("")
    lines.extend(["## 易错辨析", ""])
    for heading, content in mistakes:
        lines.extend([f"### {heading}", ""])
        lines.extend(markdown_bullets(content))
        lines.append("")
    lines.extend([
        "## 速记导图",
        "",
        f"![{number:02d} {title}知识导图](/biology/topic-{number:02d}/{map_filename})",
        "",
        "## 复习建议",
        "",
        "1. 先按考点概览口述知识框架，再对照核心知识补漏。",
        "2. 将易错辨析改写成判断题，说明错误原因和正确表述。",
        "3. 最后遮住导图关键词，独立复述概念之间的联系。",
        "",
        "---",
        "",
    ])
    nav: list[str] = []
    if previous:
        prev_num, prev_title, prev_module = previous
        prefix = "../" if prev_module != module.name else "./"
        nav.append(f"[← 上一节：{prev_num:02d} {prev_title}]({prefix}{prev_module + '/' if prev_module != module.name else ''}{prev_num:02d}-{prev_title})")
    if following:
        next_num, next_title, next_module = following
        prefix = "../" if next_module != module.name else "./"
        nav.append(f"[下一节：{next_num:02d} {next_title} →]({prefix}{next_module + '/' if next_module != module.name else ''}{next_num:02d}-{next_title})")
    lines.append(" · ".join(nav))
    lines.append("")
    output_dir = BIOLOGY_DIR / module.name
    output_dir.mkdir(parents=True, exist_ok=True)
    (output_dir / f"{number:02d}-{title}.md").write_text("\n".join(lines), encoding="utf-8")


def write_indexes(topics) -> None:
    for module in MODULES:
        selected = [topic for topic in topics if module.start <= topic[0] <= module.end]
        lines = [f"# {module.name}", "", module.description, ""]
        for number, title, points, mistakes in selected:
            lines.append(
                f"- [{number:02d} {title}](./{number:02d}-{title})：{len(points)} 个考点，{len(mistakes)} 个易错点"
            )
        lines.extend(["", "建议按编号顺序学习；复习时可直接进入薄弱专题。", ""])
        output = BIOLOGY_DIR / module.name / "index.md"
        output.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    if len(sys.argv) != 2:
        raise SystemExit("Usage: import_biology_docx.py <source-directory>")
    source = Path(sys.argv[1]).resolve()
    files = sorted(source.glob("知识清单*.docx"))
    if len(files) != 38:
        raise SystemExit(f"Expected 38 knowledge-list files, found {len(files)}")
    map_doc = next(source.glob("知识导图*.docx"), None)
    if not map_doc:
        raise SystemExit("Mind-map document not found")

    if ASSET_DIR.exists():
        shutil.rmtree(ASSET_DIR)
    topics = [parse_topic(path) for path in files]
    maps = extract_maps(map_doc)
    titles = [(topic[0], topic[1], module_for(topic[0]).name) for topic in topics]
    for index, topic in enumerate(topics):
        write_topic(
            topic,
            maps[topic[0]],
            titles[index - 1] if index else None,
            titles[index + 1] if index + 1 < len(titles) else None,
        )
    write_indexes(topics)
    print(f"Generated {len(topics)} topic pages and {len(MODULES)} module indexes.")


if __name__ == "__main__":
    main()
