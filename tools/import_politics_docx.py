"""Build the politics VitePress topic library from the supplied DOCX set.

Set POLITICS_DOCX_DIR to the directory containing the 35 source documents.
The importer groups the five document types by textbook and lesson, converts
useful tables to readable Markdown, and keeps only substantial mind-map images.
"""

from __future__ import annotations

import hashlib
import os
import re
import zipfile
from collections import OrderedDict
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
SOURCE = Path(os.environ["POLITICS_DOCX_DIR"])
POLITICS = ROOT / "政治"
ASSETS = ROOT / "public" / "politics"
IGNORED_IMAGE_HASHES: set[str] = set()
DESIRED_ASSETS: set[Path] = set()


@dataclass(frozen=True)
class Lesson:
    number: int
    title: str
    unit: str


BOOKS: OrderedDict[str, dict] = OrderedDict({
    "必修一": {
        "title": "中国特色社会主义",
        "dir": "必修一-中国特色社会主义",
        "description": "把握人类社会发展规律、中国特色社会主义的开创发展与新时代使命。",
        "lessons": [
            Lesson(1, "社会主义从空想到科学、从理论到实践的发展", "人类社会发展的进程与趋势"),
            Lesson(2, "只有社会主义才能救中国", "人类社会发展的进程与趋势"),
            Lesson(3, "只有中国特色社会主义才能发展中国", "中国特色社会主义的开创与发展"),
            Lesson(4, "只有坚持和发展中国特色社会主义才能实现中华民族伟大复兴", "中国特色社会主义的开创与发展"),
        ],
    },
    "必修二": {
        "title": "经济与社会",
        "dir": "必修二-经济与社会",
        "description": "理解基本经济制度、社会主义市场经济体制、高质量发展、收入分配与社会保障。",
        "lessons": [
            Lesson(1, "我国的生产资料所有制", "生产资料所有制与经济体制"),
            Lesson(2, "我国的社会主义市场经济体制", "生产资料所有制与经济体制"),
            Lesson(3, "我国的经济发展", "经济发展与社会进步"),
            Lesson(4, "我国的个人收入分配与社会保障", "经济发展与社会进步"),
        ],
    },
    "必修三": {
        "title": "政治与法治",
        "dir": "必修三-政治与法治",
        "description": "围绕党的领导、人民当家作主和全面依法治国，理解三者有机统一。",
        "lessons": [
            Lesson(1, "历史和人民的选择", "中国共产党的领导"),
            Lesson(2, "中国共产党的先进性", "中国共产党的领导"),
            Lesson(3, "坚持和加强党的全面领导", "中国共产党的领导"),
            Lesson(4, "人民民主专政的社会主义国家", "人民当家作主"),
            Lesson(5, "我国的根本政治制度", "人民当家作主"),
            Lesson(6, "我国的基本政治制度", "人民当家作主"),
            Lesson(7, "治国理政的基本方式", "全面依法治国"),
            Lesson(8, "法治中国建设", "全面依法治国"),
            Lesson(9, "全面依法治国的基本要求", "全面依法治国"),
        ],
    },
    "必修四": {
        "title": "哲学与文化",
        "dir": "必修四-哲学与文化",
        "description": "运用辩证唯物主义和历史唯物主义分析问题，理解文化传承与创新。",
        "lessons": [
            Lesson(1, "时代精神的精华", "探索世界与把握规律"),
            Lesson(2, "探究世界的本质", "探索世界与把握规律"),
            Lesson(3, "把握世界的规律", "探索世界与把握规律"),
            Lesson(4, "探索认识的奥秘", "认识社会与价值选择"),
            Lesson(5, "寻觅社会的真谛", "认识社会与价值选择"),
            Lesson(6, "实现人生的价值", "认识社会与价值选择"),
            Lesson(7, "继承发展中华优秀传统文化", "文化传承与文化创新"),
            Lesson(8, "学习借鉴外来文化的有益成果", "文化传承与文化创新"),
            Lesson(9, "发展中国特色社会主义文化", "文化传承与文化创新"),
        ],
    },
    "选择性必修一": {
        "title": "当代国际政治与经济",
        "dir": "选择性必修一-当代国际政治与经济",
        "description": "认识国家制度、世界多极化、经济全球化以及中国与国际组织的关系。",
        "lessons": [
            Lesson(1, "国体与政体", "各具特色的国家"),
            Lesson(2, "国家的结构形式", "各具特色的国家"),
            Lesson(3, "多极化趋势", "世界多极化"),
            Lesson(4, "和平与发展", "世界多极化"),
            Lesson(5, "中国的外交", "世界多极化"),
            Lesson(6, "走进经济全球化", "经济全球化"),
            Lesson(7, "经济全球化与中国", "经济全球化"),
            Lesson(8, "主要的国际组织", "国际组织"),
            Lesson(9, "中国与国际组织", "国际组织"),
        ],
    },
    "选择性必修二": {
        "title": "法律与生活",
        "dir": "选择性必修二-法律与生活",
        "description": "以民法典等现行法律为基础，学习民事权利、家庭、就业创业与争议解决。",
        "lessons": [
            Lesson(1, "在生活中学民法用民法", "民事权利与义务"),
            Lesson(2, "依法有效保护财产权", "民事权利与义务"),
            Lesson(3, "订约履约 诚信为本", "民事权利与义务"),
            Lesson(4, "侵权责任与权利界限", "民事权利与义务"),
            Lesson(5, "在和睦家庭中成长", "家庭与婚姻"),
            Lesson(6, "珍惜婚姻关系", "家庭与婚姻"),
            Lesson(7, "做个明白的劳动者", "就业与创业"),
            Lesson(8, "自主创业与诚信经营", "就业与创业"),
            Lesson(9, "纠纷的多元解决方式", "社会争议解决"),
            Lesson(10, "诉讼实现公平正义", "社会争议解决"),
        ],
    },
    "选择性必修三": {
        "title": "逻辑与思维",
        "dir": "选择性必修三-逻辑与思维",
        "description": "掌握逻辑思维规则、辩证思维方法与创新思维能力。",
        "lessons": [
            Lesson(1, "走进思维世界", "树立科学思维观念"),
            Lesson(2, "把握逻辑要义", "树立科学思维观念"),
            Lesson(3, "领会科学思维", "树立科学思维观念"),
            Lesson(4, "准确把握概念", "遵循逻辑思维规则"),
            Lesson(5, "正确运用判断", "遵循逻辑思维规则"),
            Lesson(6, "掌握演绎推理方法", "遵循逻辑思维规则"),
            Lesson(7, "学会归纳与类比推理", "遵循逻辑思维规则"),
            Lesson(8, "把握辩证分合", "运用辩证思维方法"),
            Lesson(9, "理解质量互变", "运用辩证思维方法"),
            Lesson(10, "推动认识发展", "运用辩证思维方法"),
            Lesson(11, "创新思维要善于联想", "提高创新思维能力"),
            Lesson(12, "创新思维要多路探索", "提高创新思维能力"),
            Lesson(13, "创新思维要力求超前", "提高创新思维能力"),
        ],
    },
})

CATEGORY_NAMES = {
    "思维导图": "知识框架与思维导图",
    "核心考点": "核心考点与辨析",
    "知识清单": "核心知识清单",
    "答题模板": "材料题答题模板",
    "易混易错": "易混易错",
}

DROP_PATTERNS = (
    "学科网", "组卷网", "未经允许", "版权所有", "侵权必究",
    "扫码", "二维码", "公众号", "微信", "获取更多",
)

TEXT_CORRECTIONS = {
    "民法典地506条": "民法典第506条",
    "因故意或者过失造成对方财产损失": "因故意或者重大过失造成对方财产损失",
    "没有敦优轨劣": "没有绝对优劣",
    "多极化化趋势": "多极化趋势",
    "掌区分2种形式": "区分2种形式",
}


def clean(text: str) -> str:
    text = text.replace("\u00a0", " ").replace("\u200b", "")
    text = text.replace("\r", " ").replace("\n", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = text.strip()
    for old, new in TEXT_CORRECTIONS.items():
        text = text.replace(old, new)
    return text


def iter_blocks(parent):
    if isinstance(parent, DocumentType):
        element = parent.element.body
    elif isinstance(parent, _Cell):
        element = parent._tc
    else:
        raise TypeError(type(parent))
    for child in element.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def identify_lesson(text: str, lessons: list[Lesson]) -> int | None:
    text = clean(text).replace("：", " ").replace(":", " ")
    if len(text) > 72:
        return None
    match = re.match(r"^第\s*([一二三四五六七八九十百0-9]+)\s*课", text)
    if not match:
        return None
    token = match.group(1)
    chinese = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5, "六": 6,
               "七": 7, "八": 8, "九": 9, "十": 10, "十一": 11,
               "十二": 12, "十三": 13}
    number = int(token) if token.isdigit() else chinese.get(token)
    if number and any(item.number == number for item in lessons):
        return number
    return None


def cell_text(cell: _Cell) -> str:
    values = []
    for paragraph in cell.paragraphs:
        value = clean(paragraph.text).replace("|", "\\|")
        if value and (not values or values[-1] != value):
            values.append(value)
    return "<br>".join(values) or " "


def normalize_rows(table: Table) -> list[list[str]]:
    rows = [[cell_text(cell) for cell in row.cells] for row in table.rows]
    if not rows:
        return []
    width = max(len(row) for row in rows)
    normalized = []
    for row in rows:
        row += [" "] * (width - len(row))
        # python-docx repeats merged-cell text; blank consecutive duplicates.
        for index in range(1, len(row)):
            if row[index] == row[index - 1] and row[index] != " ":
                row[index] = " "
        if any(value.strip() for value in row):
            normalized.append(row)
    return normalized


def table_markdown(table: Table) -> list[str]:
    rows = normalize_rows(table)
    if not rows:
        return []
    width = len(rows[0])
    longest = max(len(value) for row in rows for value in row)
    if width <= 8 and longest <= 800:
        output = [
            "| " + " | ".join(rows[0]) + " |",
            "| " + " | ".join(["---"] * width) + " |",
        ]
        output.extend("| " + " | ".join(row) + " |" for row in rows[1:])
        return output
    output = []
    headings = [re.sub(r"(?:<br>)+", "", value) for value in rows[0]]
    for row in rows[1:] if len(rows) > 1 else rows:
        fields = []
        seen_fields = set()
        for index, value in enumerate(row):
            value = re.sub(r"(?:(?<=.)<br>(?=.)){2,}", "", value)
            if not value.strip():
                continue
            label = headings[index].strip() if index < len(headings) else ""
            field = f"**{label}：**{value}" if label and label != value else value
            if field not in seen_fields:
                fields.append(field)
                seen_fields.add(field)
        if fields:
            output.append("- " + "；".join(fields))
    return output


def should_drop(text: str) -> bool:
    return not text or any(pattern in text for pattern in DROP_PATTERNS)


def paragraph_markdown(text: str) -> list[str]:
    text = clean(text)
    if should_drop(text):
        return []
    # Source titles and structural labels are redundant inside a lesson page.
    if ("知识清单" in text or "核心考点" in text or "思维导图" in text or
            "答题模板" in text or "易混易错" in text) and len(text) < 45:
        return []
    if re.match(r"^第[一二三四五六七八九十百0-9]+课", text):
        return []
    if re.match(r"^第[一二三四五六七八九十百0-9]+单元", text):
        return []
    label_match = re.match(r"^【([^】]+)】\s*(.*)$", text)
    if label_match:
        label, remainder = label_match.groups()
        if remainder:
            return [f"> **{label}：**{remainder}"]
        return [f"### {label}"]
    if re.match(r"^(核心考点|考点|知识点)\s*[一二三四五六七八九十0-9]+", text) and len(text) <= 70:
        return [f"### {text}"]
    if re.match(r"^[一二三四五六七八九十]+[、.]", text) and len(text) <= 70:
        return [f"### {text}"]
    if re.match(r"^\d+[、.]\s*[^。；]{1,45}$", text):
        return [f"#### {text}"]
    if text in {"答题要素", "答题模板", "阐释", "易错提醒", "备考策略"}:
        return [f"### {text}"]
    return [text]


def extract_images(paragraph: Paragraph, doc: DocumentType, target_dir: Path,
                   seen: set[str]) -> list[str]:
    output = []
    for blip in paragraph._p.xpath(".//a:blip"):
        rel_id = blip.get(qn("r:embed"))
        if not rel_id or rel_id not in doc.part.related_parts:
            continue
        part = doc.part.related_parts[rel_id]
        suffix = Path(str(part.partname)).suffix.lower()
        if suffix not in {".png", ".jpg", ".jpeg", ".webp"}:
            continue
        digest = hashlib.sha256(part.blob).hexdigest()
        if digest in seen or digest in IGNORED_IMAGE_HASHES:
            continue
        try:
            with Image.open(BytesIO(part.blob)) as image:
                width, height = image.size
        except Exception:
            continue
        # Keep large, legible diagrams; reject logos, banners and QR codes.
        if width < 650 or height < 240 or width * height < 350_000:
            continue
        if 0.82 <= width / height <= 1.18 and width < 1400:
            continue
        seen.add(digest)
        target_dir.mkdir(parents=True, exist_ok=True)
        filename = f"{digest[:12]}{suffix}"
        asset = target_dir / filename
        asset.write_bytes(part.blob)
        DESIRED_ASSETS.add(asset.resolve())
        output.append(filename)
    return output


def split_document(path: Path, lessons: list[Lesson], book_dir: str,
                   include_images: bool) -> dict[int, list[str]]:
    doc = Document(path)
    sections = {lesson.number: [] for lesson in lessons}
    current: int | None = None
    seen_images: set[str] = set()
    image_counts = {lesson.number: 0 for lesson in lessons}
    for block in iter_blocks(doc):
        if isinstance(block, Paragraph):
            text = clean(block.text)
            lesson_number = identify_lesson(text, lessons)
            if lesson_number is not None:
                current = lesson_number
                continue
            if current is None:
                continue
            lines = paragraph_markdown(text)
            if lines:
                sections[current].extend(lines)
                sections[current].append("")
            if include_images and image_counts[current] < 3:
                asset_dir = ASSETS / book_dir / f"lesson-{current:02d}"
                for filename in extract_images(block, doc, asset_dir, seen_images):
                    sections[current].append(
                        f"![第{current}课思维导图](/politics/{book_dir}/lesson-{current:02d}/{filename})"
                    )
                    sections[current].append("")
                    image_counts[current] += 1
        elif current is not None:
            lines = table_markdown(block)
            if lines:
                sections[current].extend(lines)
                sections[current].append("")
    return sections


def split_document_units(path: Path, lessons: list[Lesson]) -> dict[str, list[str]]:
    """Extract unit-level material as a fallback for sources without lesson headings."""
    doc = Document(path)
    unit_names = list(OrderedDict.fromkeys(lesson.unit for lesson in lessons))
    sections = {unit: [] for unit in unit_names}
    current: str | None = None
    for block in iter_blocks(doc):
        if isinstance(block, Paragraph):
            text = clean(block.text)
            matched = next((unit for unit in unit_names if unit in text and len(text) <= 55), None)
            if matched and (text.startswith("第") or text == matched):
                current = matched
                continue
            # Stop a unit fallback at a lesson heading; lesson-level extraction owns it.
            if identify_lesson(text, lessons) is not None:
                continue
            if current is not None:
                lines = paragraph_markdown(text)
                if lines:
                    sections[current].extend(lines)
                    sections[current].append("")
        elif current is not None:
            lines = table_markdown(block)
            if lines:
                sections[current].extend(lines)
                sections[current].append("")
    return sections


def locate_sources(book_key: str, title: str) -> dict[str, Path]:
    matches = [path for path in SOURCE.glob("*.docx")
               if book_key in path.name and f"《{title}》" in path.name]
    result = {}
    for category in CATEGORY_NAMES:
        candidates = [path for path in matches if f"【{category}】" in path.name]
        if len(candidates) != 1:
            raise RuntimeError(f"{book_key}《{title}》{category}: found {len(candidates)} files")
        result[category] = candidates[0]
    return result


def compact(lines: list[str]) -> list[str]:
    output = []
    previous = None
    for line in lines:
        line = line.rstrip()
        if line == previous and line:
            continue
        if not line and (not output or not output[-1]):
            continue
        output.append(line)
        previous = line
    while output and not output[-1]:
        output.pop()
    return output


def write_lesson(book_key: str, info: dict, lesson: Lesson,
                 content: dict[str, dict[int, list[str]]]) -> Path:
    target_dir = POLITICS / info["dir"]
    target_dir.mkdir(parents=True, exist_ok=True)
    target = target_dir / f"{lesson.number:02d}-{lesson.title}.md"
    lines = [
        "---",
        f"title: 第{lesson.number}课 {lesson.title}",
        "outline: [2, 3]",
        "---",
        "",
        f"# 第{lesson.number}课 {lesson.title}",
        "",
        f"> **教材位置：**{book_key}《{info['title']}》 · {lesson.unit}",
        "",
        "本页把知识清单、核心考点、答题模板、易混易错和思维导图合并为一条复习主线。",
        "",
    ]
    if book_key == "选择性必修二":
        legal_notes = {
            8: (
                "本课涉及创业与消费者保护。已按 2024 年 7 月 1 日起施行的"
                "[新修订公司法](https://www.npc.gov.cn/npc/c2/c30834/202312/t20231229_433999.html)"
                "和[消费者权益保护法实施条例](https://www.mee.gov.cn/zcwj/gwywj/202403/t20240320_1068830.shtml)核对。"
            ),
            9: (
                "本课仲裁内容已按 2026 年 3 月 1 日起施行的"
                "[新修订仲裁法](https://www.npc.gov.cn/npc/c2/c30834/202509/t20250912_447762.html)核对；"
                "劳动争议仲裁等仍适用相应专门法律。"
            ),
            10: (
                "本课诉讼程序已参考 2024 年 1 月 1 日起施行的"
                "[民事诉讼法修改决定](https://www.npc.gov.cn/npc/c2/c30834/202309/t20230901_431419.html)核对。"
            ),
        }
        note = legal_notes.get(lesson.number)
        if note:
            lines.extend([f"> **时效核对（2026 年 8 月）：**{note}", ""])
    order = ["思维导图", "核心考点", "知识清单", "答题模板", "易混易错"]
    for category in order:
        section = compact(content[category].get(lesson.number, []))
        if not section and category == "答题模板":
            section = [
                "> 原资料未按本课单列模板。作答时使用以下通用材料分析链条。",
                "",
                "1. 根据设问主体、知识限定和题型，确定本课可调用的原理。",
                "2. 概括材料中的行为、措施、原因或结果，避免整段照抄材料。",
                "3. 按“教材原理＋材料对应信息＋作用或结论”逐点展开。",
                "4. 检查每个要点是否同时回应设问和材料，删除脱离情境的套话。",
            ]
        if not section and category == "易混易错":
            section = [
                "> 原资料未按本课单列易错题。复习时重点检查概念适用范围、主体、条件以及表述中的绝对化词语。",
            ]
        if not section:
            continue
        lines.extend([f"## {CATEGORY_NAMES[category]}", ""])
        lines.extend(section)
        lines.append("")
    lines.extend([
        "## 复习建议",
        "",
        "1. 先用知识框架建立本课概念之间的联系，再脱离资料复述主干。",
        "2. 对照易混易错逐项说明“错在哪里、正确表述是什么”，避免只记结论。",
        "3. 选择一道材料题，按“材料信息—教材原理—结合分析—结论”完整作答。",
        "",
        "---",
        "",
        "> 本页依据用户提供的《2025年高考政治一轮复习知识清单（全国通用）》系列资料重编；法律、政策和时效性表述核对至 2026 年 8 月。",
        "",
    ])
    target.write_text("\n".join(compact(lines)) + "\n", encoding="utf-8")
    return target


def write_book_index(book_key: str, info: dict) -> None:
    target_dir = POLITICS / info["dir"]
    target_dir.mkdir(parents=True, exist_ok=True)
    lines = [
        "---", f"title: {book_key}《{info['title']}》", "outline: [2, 3]", "---", "",
        f"# {book_key}《{info['title']}》", "", info["description"], "",
        "## 课程目录", "",
    ]
    units: OrderedDict[str, list[Lesson]] = OrderedDict()
    for lesson in info["lessons"]:
        units.setdefault(lesson.unit, []).append(lesson)
    for unit, lessons in units.items():
        lines.extend([f"### {unit}", ""])
        for lesson in lessons:
            lines.append(f"- [第{lesson.number}课 {lesson.title}](./{lesson.number:02d}-{lesson.title})")
        lines.append("")
    lines.extend([
        "## 使用方法", "",
        "- **第一轮：**沿课程目录学习知识清单，画出概念与原理之间的联系。",
        "- **第二轮：**集中复习核心考点和易混易错，训练规范表述。",
        "- **第三轮：**使用答题模板组织材料题，但必须结合材料具体分析。",
        "",
        "## 综合探究", "",
        "综合探究不单独拆页。复习时把探究主题与相关课程串联，用真实情境检验能否综合调用教材原理。",
        "",
    ])
    if book_key == "选择性必修二":
        lines.extend([
            "## 法律时效与权威文本", "",
            "本册内容按下列现行法律法规核对至 2026 年 8 月：", "",
            "- [《中华人民共和国民法典》](https://flk.npc.gov.cn/detail?id=ff808081729d1efe01729d50b5c500bf&title=中华人民共和国民法典)（2021 年 1 月 1 日起施行）",
            "- [《中华人民共和国公司法》](https://www.npc.gov.cn/npc/c2/c30834/202312/t20231229_433999.html)（2024 年 7 月 1 日起施行）",
            "- [《中华人民共和国消费者权益保护法实施条例》](https://www.mee.gov.cn/zcwj/gwywj/202403/t20240320_1068830.shtml)（2024 年 7 月 1 日起施行）",
            "- [《中华人民共和国仲裁法》](https://www.npc.gov.cn/npc/c2/c30834/202509/t20250912_447762.html)（2026 年 3 月 1 日起施行）",
            "- [关于修改《中华人民共和国民事诉讼法》的决定](https://www.npc.gov.cn/npc/c2/c30834/202309/t20230901_431419.html)（2024 年 1 月 1 日起施行）",
            "",
            "> 学习资料用于理解教材原理，不构成针对具体案件的法律意见。",
            "",
        ])
    (target_dir / "index.md").write_text("\n".join(lines), encoding="utf-8")


def write_subject_home() -> None:
    lines = [
        "---", "title: 政治", "outline: [2, 3]", "---", "", "# 政治", "",
        "政治复习不能停留在背诵结论。要把教材原理、材料信息和规范表达连接起来，形成“读材料—定范围—找原理—作分析”的完整能力。",
        "", "## 教材专题库", "",
    ]
    for book_key, info in BOOKS.items():
        lines.extend([
            f"### [{book_key}《{info['title']}》](./{info['dir']}/)", "",
            info["description"], "",
        ])
    lines.extend([
        "## 三轮复习路径", "",
        "1. **通读建网：**按课学习知识框架和清单，能口头解释核心概念。",
        "2. **辨析固基：**集中处理易混易错，用完整教材语言订正错误观点。",
        "3. **材料迁移：**依据设问限定知识范围，用材料关键词调用原理并逐点分析。",
        "", "## 内容说明", "",
        "- 专题内容由 35 份复习资料按课重编，重复表述已合并，明显课名错误已订正。",
        "- 答题模板用于提示分析角度，不应脱离设问和材料机械照搬。",
        "- 法律、政策和国际形势会变化；相关内容已核对至 2026 年 8 月，使用时仍应关注最新权威发布。",
        "- 站内资料用于复习归纳，不能替代教材原文、现行法律文本或学校教学安排。",
        "",
    ])
    (POLITICS / "前言.md").write_text("\n".join(lines), encoding="utf-8")


def remove_stale_outputs(expected: set[Path]) -> None:
    for path in POLITICS.rglob("*.md"):
        if path.name == "前言.md":
            continue
        if path.resolve() not in expected:
            path.unlink()
    for directory in sorted(POLITICS.rglob("*"), reverse=True):
        if directory.is_dir() and not any(directory.iterdir()):
            directory.rmdir()


def main() -> None:
    source_sets = {
        book_key: locate_sources(book_key, info["title"])
        for book_key, info in BOOKS.items()
    }
    media_books: dict[str, set[str]] = {}
    for book_key, sources in source_sets.items():
        with zipfile.ZipFile(sources["思维导图"]) as archive:
            for name in archive.namelist():
                if not name.startswith("word/media/"):
                    continue
                digest = hashlib.sha256(archive.read(name)).hexdigest()
                media_books.setdefault(digest, set()).add(book_key)
    # Repeated cross-book images are promotional banners/QR codes, not diagrams.
    IGNORED_IMAGE_HASHES.update(
        digest for digest, books in media_books.items() if len(books) >= 2
    )
    expected: set[Path] = set()
    for book_key, info in BOOKS.items():
        sources = source_sets[book_key]
        content = {}
        for category, source in sources.items():
            content[category] = split_document(
                source, info["lessons"], info["dir"], category == "思维导图"
            )
            unit_fallbacks = split_document_units(source, info["lessons"])
            for lesson in info["lessons"]:
                if not compact(content[category][lesson.number]):
                    fallback = compact(unit_fallbacks.get(lesson.unit, []))
                    if fallback:
                        content[category][lesson.number] = [
                            f"> 本节源资料按“{lesson.unit}”单元汇总，以下内容供本课关联复习。",
                            "",
                            *fallback,
                        ]
        for lesson in info["lessons"]:
            target = write_lesson(book_key, info, lesson, content)
            expected.add(target.resolve())
            print(target.relative_to(ROOT))
        write_book_index(book_key, info)
        expected.add((POLITICS / info["dir"] / "index.md").resolve())
    write_subject_home()
    remove_stale_outputs(expected)
    if ASSETS.exists():
        for path in ASSETS.rglob("*"):
            if path.is_file() and path.resolve() not in DESIRED_ASSETS:
                path.unlink()
        for directory in sorted(ASSETS.rglob("*"), reverse=True):
            if directory.is_dir() and not any(directory.iterdir()):
                directory.rmdir()


if __name__ == "__main__":
    main()
